"""Khao sat cau truc file SO TAY SINH VIEN 2024-2025.docx.

In ra:
1. Thong ke tong quan (so doan, bang)
2. Cac heading / doan co dang "Chuong", "Dieu", "Phan" de hieu cau truc
3. Mau 30 doan dau tien de xem cach trinh bay
"""
import re
from docx import Document
from collections import Counter

PATH = r"data\raw\SO TAY SINH VIEN 2024-2025.docx"
doc = Document(PATH)

paras = doc.paragraphs
print(f"Tong so paragraph: {len(paras)}")
print(f"Tong so bang: {len(doc.tables)}")
print()

# --- 1. Dem cac style duoc dung ---
styles = Counter(p.style.name for p in paras if p.text.strip())
print("=== Styles duoc dung (top 15) ===")
for s, c in styles.most_common(15):
    print(f"  {c:5d}  {s}")
print()

# --- 2. Tim cac doan co dau hieu cau truc phap ly ---
patterns = {
    "Chuong": re.compile(r"^\s*(ch[uư][oơ]*ng)\s+[IVX0-9]+", re.I),
    "Dieu":   re.compile(r"^\s*[Đđ](i|і)ều\s+\d+", re.I),
    "Phan":   re.compile(r"^\s*(phần|phan)\s+[IVX0-9]+", re.I),
    "Muc":    re.compile(r"^\s*(mục|muc)\s+\d+", re.I),
}
found = {k: [] for k in patterns}
for i, p in enumerate(paras):
    t = p.text.strip()
    if not t:
        continue
    for name, rx in patterns.items():
        if rx.match(t):
            found[name].append((i, t[:80], p.style.name))

print("=== Dem dau hieu cau truc ===")
for name, items in found.items():
    print(f"  {name}: {len(items)}")
print()

# --- 3. In 20 "Chuong" dau ---
print("=== Cac 'Chuong' dau tien ===")
for i, t, st in found["Chuong"][:20]:
    print(f"  [para {i}] ({st}) {t}")
print()

# --- 4. In 25 "Dieu" dau ---
print("=== Cac 'Dieu' dau tien ===")
for i, t, st in found["Dieu"][:25]:
    print(f"  [para {i}] ({st}) {t}")
print()

# --- 5. Mau 25 doan khac trong dau tien co noi dung ---
print("=== 25 doan co noi dung dau tien ===")
count = 0
for i, p in enumerate(paras):
    t = p.text.strip()
    if not t:
        continue
    print(f"  [para {i}] ({p.style.name}) {t[:100]}")
    count += 1
    if count >= 25:
        break
